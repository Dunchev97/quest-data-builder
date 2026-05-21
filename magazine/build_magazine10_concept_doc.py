from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "Magazine10_concept_DetectiveCats.docx"


ACCENT = "1F4E78"
SOFT_BLUE = "DDEBF7"
SOFT_GREEN = "E2F0D9"
SOFT_YELLOW = "FFF2CC"
SOFT_RED = "F4CCCC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        set_cell_shading(header_cells[idx], SOFT_BLUE)
        if widths:
            header_cells[idx].width = Cm(widths[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                cells[idx].width = Cm(widths[idx])


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    add_table(doc, ["Параметр", "Значение"], [[left, right] for left, right in rows], widths=[5.5, 11.5])


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(4)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)
    return doc


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Домовита №10\n"Детективное агентство котиков"')
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Концепт-документ журнала для контента, арта и разработки")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    add_key_value_table(
        doc,
        [
            ("Выпуск", "Домовита, номер 10"),
            ("Тема выпуска", "Уютный сказочный детектив, котики-сыщики, домашние улики и добрые приключения"),
            ("Главный приз", 'Фасад для дома "Детективное агентство котиков"'),
            ("Большая цепочка", "12 подвигов героя с финальным питомцем: лисенок Луи"),
            ("Ключевой объект DIY", 'Фотозона "Детективное агентство котиков"'),
            ("Важно по названию", 'В публичных текстах использовать "агентство"; форма "агентсво" считается опечаткой.'),
        ],
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Документ собран на основе Magazine9.proto.js, quest-прототипов Magazine9 и вводных по выпуску 10.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_overview(doc: Document) -> None:
    add_heading(doc, "1. Короткая идея", 1)
    doc.add_paragraph(
        'Выпуск "Детективное агентство котиков" превращает обычный журнал в теплое расследование без мрачности: '
        "домовенок листает рубрики, собирает новые знания и награды, помогает собрать фотозону для маленьких сыщиков, "
        "а в финале получает фасад агентства. Детективная тема держится на уютных уликах: лапки, лупы, клубки, записки, "
        "карта с ниточками и важный котик, который делает вид, будто все раскрыл еще утром."
    )
    doc.add_paragraph(
        "Журнал не становится криминальной историей. Это добрый бытовой хаос в стиле Домовят: пропали клубки, перепутались "
        "папки, Кощей проверяет надежность тайников, а соседушки помогают довести дело до счастливого финала."
    )

    add_heading(doc, "2. Цели выпуска", 1)
    add_bullets(
        doc,
        [
            "Собрать выпуск №10 на базе проверенной структуры Magazine9, но без удаленных страниц color_furniture, color_clothes, ДомоВестник и test.",
            'Сделать тему "Детективное агентство котиков" главной рамкой для DIY, search и суперприза.',
            "Ввести новые энциклопедические сущности: алмаз, психотрию возвышенную, кизил и окапи.",
            "Трижды использовать clothes-шаблон для трех пар костюмов: скандинавский, ярмарочный, подводное ателье.",
            'Заменить "Крохо-Гороскоп" на механику "12 подвигов героя" с финальным питомцем лисенком Луи.',
            "Дать разработчикам ясную карту страниц, quest-групп, hidden-флагов, rewards и мест, где нужны реальные classnames/id.",
        ],
    )

    add_heading(doc, "3. Тон и визуальное направление", 1)
    add_table(
        doc,
        ["Блок", "Решение"],
        [
            ["Настроение", "Уютный сказочный детектив: любопытство, мягкая ирония, домашняя тайна, котики-сыщики."],
            ["Не использовать", "Мрачный noir, угрозы, настоящие преступления, формальные фразы вроде 'необходимо' и 'требуется'."],
            ["Визуальные мотивы", "Лупа, следы лап, клубки, папки с делами, карта с ниточками, теплый свет лампы, табличка агентства."],
            ["Цвета", "Чернильно-синий, теплый янтарный, молочный бумажный, зеленый садовый, клюквенный акцент."],
            ["Главная фантазия игрока", "Домовенок становится помощником доброго кото-детективного бюро и постепенно обустраивает его у себя дома."],
        ],
        widths=[4.0, 12.5],
    )


def add_structure(doc: Document) -> None:
    add_heading(doc, "4. Структура страниц", 1)
    rows = [
        ["0", "cover", "", "Обложка выпуска 10", "Magazine10Cover"],
        ["1", "content", "Содержание", "Оглавление, часть 1", "Magazine10Content1"],
        ["2", "content", "Содержание", "Оглавление, часть 2", "Magazine10Content2"],
        ["3", "stone", "Каменная сказка", "Алмаз", "Magazine10Stone"],
        ["4", "flower", "Лютики-цветочки", "Цветочные губы, психотрия возвышенная", "Magazine10Flower"],
        ["5", "tree", "В тени дерев", "Кизил", "Magazine10Tree"],
        ["6", "achieve", "Уши, лапы, хвост", "Окапи через achieve-шаблон", "Magazine10Animal"],
        ["7", "photo", "Все о фотографии!", "Пока копия страницы фото из выпуска 9", "Magazine10Photo"],
        ["8", "clothes", "Шитье и рукоделие", "Костюм в скандинавском стиле: мальчик + девочка", "Magazine10ClothesScandiBoy/Girl"],
        ["9", "clothes", "Шитье и рукоделие", "Костюм ярмарочный: мальчик + девочка", "Magazine10ClothesFairBoy/Girl"],
        ["10", "clothes", "Шитье и рукоделие", "Костюм в стиле подводного ателье: мальчик + девочка", "Magazine10ClothesAtelierBoy/Girl"],
        ["11", "diy", "Советы Самоделкина", 'Фотозона "Детективное агентство котиков", 3 этапа', "Magazine10Diy"],
        ["12", "interior_title", "Интерьер", "Вступление к BOHO-интерьеру", ""],
        ["13", "interior_shop", "Интерьер", "12 декоров в стиле BOHO", ""],
        ["14", "beauty_hair", "Коса - девичья краса", "4 женские прически", ""],
        ["15", "beauty_hair", "Причешись-ка!", "4 мужские прически", ""],
        ["16", "beauty_makeup", "Уроки преображения", "3 образа + Забавный стилист", ""],
        ["17", "cook", "Пальчики оближешь!", "Запеченный ролл-пирог", ""],
        ["18", "interview", "В гостях...", "Интервью с Кощеем Бессмертным", ""],
        ["19", "achieve", "Профессия - Кукольный модельер", "Профессия через achieve-шаблон", "Magazine10Achieve"],
        ["20", "search", "Расследование начинается!", "Пока копия поисковой цепочки", "Magazine10Search"],
        ["21", "tiny_horoscope", "12 подвигов героя", "12 заданий и финальный питомец лисенок Луи", "Magazine10HeroFeats"],
        ["22", "surprise", "Суперприз", 'Фасад "Детективное агентство котиков"', "Magazine10"],
    ]
    add_table(doc, ["№", "template", "label_title", "Контент", "quest_group"], rows, widths=[1.0, 3.2, 4.0, 6.2, 4.0])

    add_heading(doc, "5. Что убираем из выпуска 9", 1)
    add_table(
        doc,
        ["Было", "Решение"],
        [
            ["color_furniture / Квартирный вопрос", "Полностью убрать."],
            ["color_clothes / Модный приговор", "Полностью убрать обе страницы."],
            ["ДомоВестник", "Полностью убрать обе страницы."],
            ["test / Тест", "Полностью убрать."],
            ["animal / Уши, лапы, хвост", "Не использовать animal template; окапи сделать через achieve."],
            ["Крохо-Гороскоп", 'Переименовать и переписать в "12 подвигов героя", механику можно оставить tiny_horoscope.'],
        ],
        widths=[6.0, 10.5],
    )


def add_page_concepts(doc: Document) -> None:
    add_heading(doc, "6. Концепты ключевых рубрик", 1)
    rows = [
        ["Алмаз", "Камень ясности, блеска и крепкой воли. Текст строится вокруг солнечных искр, твердой мысли и смелого сердца."],
        ["Цветочные губы", "Психотрия возвышенная выглядит как смешная записка природы: не хмурься, дружок, даже листья умеют улыбаться."],
        ["Кизил", "Садовый лекарь с рубиновыми ягодками. Ассоциации: запасливость, варенье, теплый сад."],
        ["Окапи", "Лесной тихоня в полосатых чулочках. Страница выглядит как achieve, но по содержанию заменяет старую animal-рубрику."],
        ["Три костюма", "Скандинавский уют, ярмарочная праздничность, подводное ателье с перламутром и пуговицами-ракушками."],
        ["DIY", "Фотозона, где домовенок выглядит главным следователем по делу пропавших клубков."],
        ["Интерьер BOHO", "12 предметов: лен, ротанг, макраме, песочные узоры, ловец снов, домашние травы."],
        ["Beauty", "Женские и мужские прически, макияж/борода/усы, плюс Забавный стилист - хомяк барбер-стилист."],
        ["Кулинария", "Запеченный ролл-пирог: нарезается как ролл, подается как пирог, выглядит как домашнее чудо."],
        ["Интервью", "Кощей Бессмертный добрый, тщеславный и хозяйственный: не теряет ключи, а проверяет надежность тайников."],
        ["Search", "Пока техническая копия Magazine9Search. По желанию можно тематически заменить улики под кото-детективную историю."],
        ["Суперприз", 'Фасад дома "Детективное агентство котиков": усатый адрес для самых важных домашних расследований.'],
    ]
    add_table(doc, ["Рубрика", "Концепт"], rows, widths=[4.2, 12.3])

    add_heading(doc, "7. Текстовые опорные фразы", 1)
    add_table(
        doc,
        ["Где", "Текст"],
        [
            ["DIY", "Мастерим фотозону, где каждый домовенок сможет выглядеть как главный следователь по делу пропавших клубков."],
            ["Суперприз", "Теперь у самого серьезного дела будет самый усатый адрес."],
            ["Суперприз", "За этим фасадом легко представить шуршащие папки, карту с ниточками и важного котика, который делает вид, что все раскрыл еще утром."],
            ["12 подвигов", "Выполни 12 добрых подвигов героя и собери все памятные значки. Когда все дела будут закончены, лисенок Луи отправится с тобой в новые приключения."],
        ],
        widths=[3.5, 13.0],
    )


def add_quests(doc: Document) -> None:
    add_heading(doc, "8. Quest-логика и hidden-флаги", 1)
    doc.add_paragraph(
        "Root quest `Magazine10_1` копирует роль `Magazine9_1`: он проверяет, что игрок закрыл игровые страницы журнала, "
        "а затем дает забрать главный surprise. Ниже список проверок, которые должны заменить старые Magazine9-флаги."
    )
    rows = [
        ["3", "Magazine10_Stone_Hidden", "Magazine10Stone_1"],
        ["4", "Magazine10_Flower_Hidden", "Magazine10Flower_1"],
        ["5", "Magazine10_Tree_Hidden", "Magazine10Tree_1"],
        ["6", "Magazine10_Animal_Hidden", "Magazine10Animal_1"],
        ["7", "Magazine10_Photo_Hidden", "Magazine10Photo_1"],
        ["8", "Magazine10_ClothesScandi_Hidden", "Magazine10ClothesScandiBoy/Girl_1"],
        ["9", "Magazine10_ClothesFair_Hidden", "Magazine10ClothesFairBoy/Girl_1"],
        ["10", "Magazine10_ClothesAtelier_Hidden", "Magazine10ClothesAtelierBoy/Girl_1"],
        ["11", "Magazine10_Diy_Hidden", "Magazine10Diy_3"],
        ["19", "Magazine10_Achieve_Hidden", "Magazine10Achieve_1"],
        ["20", "Magazine10_Search_Hidden", "Magazine10Search_5"],
        ["21", "PetFoxLouie или Magazine10_HeroFeats_Hidden", "Magazine10HeroFeats_final"],
        ["22", "action Magazine10 / take_magazine_surprise", "surprise page"],
    ]
    add_table(doc, ["Страница", "Что проверяет root quest", "Кто ставит"], rows, widths=[2.2, 7.2, 6.5])

    add_heading(doc, "9. 12 подвигов героя", 1)
    hero_rows = [
        ["1", "Пекарь", "Испечь угощение для дороги."],
        ["2", "Лесник-следопыт", "Найти верную тропинку и не растерять шишки."],
        ["3", "Мастер игрушек", "Починить игрушку, которая поднимает настроение."],
        ["4", "Плотник", "Сколотить крепкую полку для приключений."],
        ["5", "Собиратель", "Собрать полезные мелочи без лишней суеты."],
        ["6", "Хранитель сада", "Ухаживать за садом и защитить ростки."],
        ["7", "Картограф", "Нарисовать карту, где даже поворот за печкой отмечен."],
        ["8", "Портной", "Сшить походный наряд с тайным карманом."],
        ["9", "Фонарщик", "Зажечь огоньки, чтобы никто не заблудился."],
        ["10", "Смотритель маяка", "Подать световой знак с самой высокой точки."],
        ["11", "Охотник за сокровищами", "Найти блестящую находку и не зазнаться."],
        ["12", "Почтальон", "Доставить важную весточку точно в срок."],
        ["Final", "Лисенок Луи", "Финальная награда за все 12 подвигов."],
    ]
    add_table(doc, ["№", "Роль", "Смысл задания"], hero_rows, widths=[1.5, 5.0, 10.0])


def add_interview(doc: Document) -> None:
    add_heading(doc, "10. Мини-интервью с Кощеем", 1)
    dialogue = [
        ("Д", 'Кощей Бессмертный, спасибо, что заглянули в "Домовиту". Говорят, вы храните столько сундуков, что сами иногда теряете ключи. Это правда?'),
        ("Кощей", "Клевета с легким звоном правды. Ключи я не теряю, я их прячу в надежные места, а потом проверяю, насколько места надежные. Иногда проверка длится триста лет, но порядок есть порядок."),
        ("Д", "А почему вас все считают таким суровым?"),
        ("Кощей", 'Вид у меня деловой. Если стоять у окна в плаще и задумчиво считать монеты, соседи сразу решают: "Ох, опять что-то замышляет". А я, может, пирог остужаю и думаю, хватит ли корицы.'),
        ("Д", "Пирог? Неожиданно!"),
        ("Кощей", "Бессмертие длинное, дружок. За это время и пироги научишься печь, и носки штопать, и спорить с чайником. Самое трудное - не колдовство, а достать из духовки противень и не забыть рукавички."),
        ("Д", "Что бы вы пожелали нашим читателям?"),
        ("Кощей", "Берегите свои иголочки, ключики и добрые слова. Первые две вещи помогают найти сокровища, а третья - тех, с кем сокровищами приятно делиться."),
    ]
    add_table(doc, ["Кто", "Реплика"], [[a, b] for a, b in dialogue], widths=[2.5, 14.0])


def add_assets_and_todo(doc: Document) -> None:
    add_heading(doc, "11. Основные новые ассеты", 1)
    rows = [
        ["Камень", "Алмаз", "Нужны piles/help asset/иконки."],
        ["Цветок", "Цветочные губы / психотрия возвышенная", "Нужны collection_id, recipe_id, seed, condition."],
        ["Дерево", "Кизил", "Reward tree asset."],
        ["Животное", "Окапи", "Визуально через achieve page."],
        ["DIY", 'Фотозона "Детективное агентство котиков"', "Награда за 3-й шаг DIY."],
        ["Интерьер", "12 BOHO-декоров", "Проверить вместимость одной interior_shop page."],
        ["Clothes", "3 пары костюмов для мальчика/девочки", "Скандинавский, ярмарочный, подводное ателье."],
        ["Beauty", "8 причесок, 3 образа, Забавный стилист", "Нужны classnames и payment_service_id."],
        ["Cook", "Запеченный ролл-пирог", "Новый food reward."],
        ["HeroFeats", "12 значков + питомец лисенок Луи", "Pet classname пока placeholder."],
        ["Surprise", 'Фасад "Детективное агентство котиков"', "Главный приз журнала."],
    ]
    add_table(doc, ["Категория", "Ассет", "Комментарий"], rows, widths=[3.5, 6.0, 7.0])

    add_heading(doc, "12. Открытые вопросы для разработки", 1)
    add_numbered(
        doc,
        [
            "Выдать уникальные numeric id для всех новых QuestPrototype и task identifier вида e....",
            "Подтвердить реальные classnames всех новых ассетов, включая питомца Луи и фасад.",
            "Уточнить flower-набор для психотрии: collection_id, recipe_id, asset_seed, asset_for_sale.",
            "Проверить, что template achieve корректно отображает окапи на странице 'Уши, лапы, хвост'.",
            "Проверить, помещает ли interior_shop 12 BOHO-декоров на одну страницу.",
            "Уточнить payment_service_id для beauty_makeup и Забавного стилиста.",
            "Решить, остаются ли photo/search полностью копией Magazine9 или получают новые объекты/улики.",
            "Заполнить economy: amount, price, rewards, тайминги и стоимость покупок.",
        ],
    )


def add_appendix(doc: Document) -> None:
    add_heading(doc, "13. Связанные рабочие файлы", 1)
    add_table(
        doc,
        ["Файл", "Назначение"],
        [
            ["magazine/Magazine10_GDD.md", "Текстовая ГДД-карта выпуска."],
            ["magazine/Magazine10_dev_table.xlsx", "Dev-таблица по страницам, quest-файлам, задачам, ассетам и TODO."],
            ["magazine/build_magazine10_dev_table.py", "Генератор dev-таблицы."],
            ["magazine/build_magazine10_concept_doc.py", "Генератор этого концепт-документа."],
        ],
        widths=[7.0, 9.5],
    )


def build() -> None:
    doc = setup_document()
    add_title_page(doc)
    add_overview(doc)
    add_structure(doc)
    add_page_concepts(doc)
    add_quests(doc)
    add_interview(doc)
    add_assets_and_todo(doc)
    add_appendix(doc)
    doc.save(OUT_PATH)


def smoke_test() -> None:
    doc = Document(OUT_PATH)
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    all_text = text + "\n" + table_text
    for marker in ["????", "Рџ", "Рє", "Рё"]:
        if marker in all_text:
            raise AssertionError(f"Encoding marker found: {marker}")
    for required in ["Детективное агентство котиков", "Окапи", "Кощей", "Кукольный модельер", "лисенок Луи"]:
        if required not in all_text:
            raise AssertionError(f"Missing required text: {required}")
    print(f"saved: {OUT_PATH}")
    print(f"paragraphs: {len(doc.paragraphs)}")
    print(f"tables: {len(doc.tables)}")


if __name__ == "__main__":
    build()
    smoke_test()
